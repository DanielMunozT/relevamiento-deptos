import openpyxl # pip3.5 install openpyxl
import docx # pip3.5 install python-docx
import re
import sys
import os.path
import tkinter as tk
from tkinter import filedialog

def esCeldaTitulo(celda):
    return celda.font.b

def esCeldaDiccionario(celdaStr):
    return '=' in celdaStr

# Completa el diccionario 'dic' si es una fila diccionario, y de serlo retorna True
def esFilaDiccionario(hoja, fila, dic):
    eraDic = False
    for col in range(1, hoja.max_column + 1):
        celdaStr = str(hoja.cell(row=fila, column=col).value)
        if esCeldaDiccionario(celdaStr):
            clave, valor = celdaStr.split('=')
            if col not in dic:
                dic[col] = {}
            dic[col][clave] = valor
            eraDic = True
    return eraDic

# Completa el diccionario 'titulo' si es una fila título, y de serlo retorna True
def esFilaTitulo(hoja, fila, titulo):
    eraTitulo = False
    for col in range(1, hoja.max_column + 1):
        celda = hoja.cell(row=fila, column=col)
        if esCeldaTitulo(celda):
            titulo[col] = celda.value
            eraTitulo = True
    return eraTitulo

# Se fija si toda la fila está vacía
def esFilaVacia(hoja, fila):
    for col in range(1, hoja.max_column + 1):
        celda = hoja.cell(row=fila, column=col)
        if celda.value:
            return False
    return True

# Si la 'claveStr' se encuentra en el diccionario retorna la traducción, sino retorna la misma 'claveStr'
def traducirConDic(claveStr, dic):
    if claveStr in dic:
        return dic[claveStr]
    else:
        return claveStr

# Obtiene el rango especificado en una celda
def rangoCelda(celda, dic):
    if not celda.value:
        return None
    celdaStr = str(celda.value)
    if ',' in celdaStr:
        rangos = [x.strip() for x in celdaStr.split(',')]
    else:
        rangos = [celdaStr.strip()]
    elementos = []
    for rango in rangos:
        if '-' in rango:
            inicio, fin = rango.split('-')
            inicioStrip = inicio.strip()
            finStrip = fin.strip()
            if inicioStrip.isalpha() and finStrip.isalpha() and len(inicioStrip) == 1 and len(finStrip) == 1:
                for e in range(ord(inicioStrip), ord(finStrip) + 1):
                    elementos += [traducirConDic(chr(e), dic)]
            elif inicioStrip.isdigit():
                for e in range(int(inicioStrip), int(finStrip) + 1):
                    elementos += [traducirConDic(str(e), dic)]
            else:
                print('No entiendo el rango: ', rango, ' en la celda: ', celdaStr)
        else:
            elementos += [traducirConDic(str(rango), dic)]
    return elementos

# Completa los primeros elementos en blanco de un rango con los mismo elementos de algún rango anterior
def completarConElAnterior(rango, rangoAnterior, maxRango):
    for i in range(1, maxRango + 1):
        if i in rango: # Al primero retorno
            return rango
        else:
            if i in rangoAnterior:
                rango[i] = rangoAnterior[i]
    return rango

# Expande una línea
def expandirClave(listaDic, clave, valores):
    nuevaListaDic = []
    if len(listaDic) > 0:
        for dic in listaDic:
            for valor in valores:
                nuevoDic = dic.copy()
                nuevoDic[clave] = valor
                nuevaListaDic += [nuevoDic]
    else:
        for valor in valores:
            nuevaListaDic += [{clave: valor}]
    return nuevaListaDic

# Escribe la expansión de una línea y retorna el número de líneas escritas
def escribirLineasPlanilla(hojaOut, listaDic, desdeFila):
    fila = desdeFila
    for filaDic in listaDic:
        for col in filaDic:
            celda = hojaOut.cell(row=fila, column=col)
            celda.value = filaDic[col]
        fila += 1
    return fila - desdeFila

# Copia (solo texto) todos los párrafos del documento original al nuevo documento
# reemplazando las plantillas y gagregando un salto de página al final
def agregarParrafoWord(parrafos, docxOut, dic, invTitulo):
    for parrafo in parrafos:
        newText = parrafo.text
        for match in agregarParrafoWord.regex.findall(parrafo.text):
            if match[2]:
                tag = match[2]
                if tag in invTitulo and invTitulo[tag] in dic:
                    newText = newText.replace(match[0], match[1] + dic[invTitulo[tag]] + match[3])
                else:
                    newText = newText.replace(match[0], '')
            else:
                tag = match[4]
                if tag in invTitulo and invTitulo[tag] in dic:
                    newText = newText.replace(match[0], dic[invTitulo[tag]])
                else:
                    newText = newText.replace(match[0], '')
        docxOut.add_paragraph(newText)
    docxOut.add_page_break()

agregarParrafoWord.regex = re.compile(r'''(
    {([^{}<>]*?)<(\w+)>([^{}<>]*?)}
    |
    <(\w+)>
)''', re.VERBOSE)


# Procesa una hoja
def procesarHoja(hoja, hojaOut, parrafos, docxOut):
    titulo = {}
    invTitulo = {}
    dic = {}
    rangoAnterior = {}
    filaOut = 1
    for fila in range(1, hoja.max_row + 1):
        if esFilaDiccionario(hoja, fila, dic):
            invTitulo = {v: k for k, v in titulo.items()}
        elif esFilaVacia(hoja, fila):
            filaOut += 1
        elif esFilaTitulo(hoja, fila, titulo):
            continue
        else:
            rangos = {}
            for col in range(1, hoja.max_column + 1):
                celda = hoja.cell(row=fila, column=col)
                if col not in dic:
                    dic[col] = {}
                rango = rangoCelda(celda, dic[col])
                if rango:
                    rangos[col] = rango
            if rangoAnterior and len(rangoAnterior) > 0:
                rangos = completarConElAnterior(rangos, rangoAnterior, hoja.max_column)
            rangoAnterior = rangos
            listaDic = []
            for col in list(rangos):
                listaDic = expandirClave(listaDic, col, rangos[col])
            filaOut += escribirLineasPlanilla(hojaOut, listaDic, filaOut)
            for dicElem in listaDic:
                agregarParrafoWord(parrafos, docxOut, dicElem, invTitulo)

# Procesa el documento
def procesarDocumento(xlsxIn, xlsxOut, docxIn, docxOut):
    xlsxOut.remove_sheet(xlsxOut.active)
    parrafos = docxIn.paragraphs
    for hoja in xlsxIn.worksheets:
        hojaOut = xlsxOut.create_sheet(title=hoja.title);
        procesarHoja(hoja, hojaOut, parrafos, docxOut)

# Se puede (1) pasar el nombre por línea de comando, (2) usar un nombre por default o (3) abrir con un diálogo
def nombreExcelInput():
    for arg in sys.argv[1:]:
        name = nombreExcelInput.xlsxRegex.search(arg)
        if name:
            return name
    if os.path.isfile(nombreExcelInput.defaultName):
        return nombreExcelInput.defaultName
    root = tk.Tk()
    root.withdraw()
    return filedialog.askopenfilename(filetypes=[('Documentos de Excel', '*.xlsx')], message='Relevamiento en Excel')

nombreExcelInput.xlsxRegex = re.compile(r'^(.*?\.xlsx)$')
nombreExcelInput.defaultName = 'relevamiento.xlsx'


# Se puede (1) pasar el nombre por línea de comando, (2) usar un nombre por default o (3) abrir con un diálogo
def nombreWordInput():
    for arg in sys.argv[1:]:
        name = nombreWordInput.docxRegex.search(arg)
        if name:
            return name
    if os.path.isfile(nombreWordInput.defaultName):
        return nombreWordInput.defaultName
    root = tk.Tk()
    root.withdraw()
    return filedialog.askopenfilename(filetypes=[('Documentos de Word', '*.docx')], message='Plantilla en Word')

nombreWordInput.docxRegex = re.compile(r'^(.*?\.docx)$')
nombreWordInput.defaultName = 'plantilla.docx'

def procesarDocumentos():
    nombreExcel = nombreExcelInput()
    if nombreExcel:
        nombreWord = nombreWordInput()
        if nombreWord:
            xlsxIn = openpyxl.load_workbook(nombreExcel)
            xlsxOut = openpyxl.Workbook()
            docxIn = docx.Document(nombreWord)
            docxOut = docx.Document()
            procesarDocumento(xlsxIn, xlsxOut, docxIn, docxOut)
            xlsxOut.save('relevamiento.listo.xlsx')
            docxOut.save('plantilla.lista.docx')

procesarDocumentos()